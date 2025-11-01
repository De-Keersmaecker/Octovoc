"""Service for generating PDF quotes from Word templates"""
import os
import tempfile
from datetime import datetime
from docx import Document
from docx2pdf import convert


def generate_quote_pdf(name, school_name, billing_address, num_classrooms, num_students, num_teacher_accounts):
    """
    Generate a PDF quote from the Word template

    Returns: path to generated PDF file
    """
    # Get template path
    template_path = os.path.join(
        os.path.dirname(os.path.dirname(__file__)),
        'templates',
        'offerte_template.docx'
    )

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    # Calculate total price (€0.95 per student)
    total_price = num_students * 0.95

    # Format billing address or use school name if not provided
    formatted_address = billing_address if billing_address else school_name

    # Load template
    doc = Document(template_path)

    # Replace placeholders in all paragraphs
    for paragraph in doc.paragraphs:
        if '[aantal klascodes]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[aantal klascodes]', str(num_classrooms))
        if '[aantal leerlingen]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[aantal leerlingen]', str(num_students))
        if '[aantal lerarencodes]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[aantal lerarencodes]', str(num_teacher_accounts))
        if '[facturatie-adres]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[facturatie-adres]', formatted_address)
        if '[totaal]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[totaal]', f"€ {total_price:.2f}")

    # Replace placeholders in all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '[aantal klascodes]' in paragraph.text:
                        paragraph.text = paragraph.text.replace('[aantal klascodes]', str(num_classrooms))
                    if '[aantal leerlingen]' in paragraph.text:
                        paragraph.text = paragraph.text.replace('[aantal leerlingen]', str(num_students))
                    if '[aantal lerarencodes]' in paragraph.text:
                        paragraph.text = paragraph.text.replace('[aantal lerarencodes]', str(num_teacher_accounts))
                    if '[facturatie-adres]' in paragraph.text:
                        paragraph.text = paragraph.text.replace('[facturatie-adres]', formatted_address)
                    if '[totaal]' in paragraph.text:
                        paragraph.text = paragraph.text.replace('[totaal]', f"€ {total_price:.2f}")

    # Create temporary directory for output
    temp_dir = tempfile.mkdtemp()

    # Save modified Word document
    word_output = os.path.join(temp_dir, 'offerte.docx')
    doc.save(word_output)

    # Convert to PDF
    pdf_output = os.path.join(temp_dir, f'Offerte_Octovoc_{school_name.replace(" ", "_")}.pdf')

    try:
        # Use LibreOffice for conversion (installed on most Linux systems)
        import subprocess
        result = subprocess.run([
            'libreoffice',
            '--headless',
            '--convert-to',
            'pdf',
            '--outdir',
            temp_dir,
            word_output
        ], capture_output=True, text=True, timeout=30)

        # Rename output file to desired name
        generated_pdf = os.path.join(temp_dir, 'offerte.pdf')
        if os.path.exists(generated_pdf):
            os.rename(generated_pdf, pdf_output)

        if not os.path.exists(pdf_output):
            raise Exception(f"PDF conversion failed: {result.stderr}")

    except FileNotFoundError:
        # LibreOffice not found, try alternative method
        raise Exception("LibreOffice not installed - PDF conversion not available")

    return pdf_output
